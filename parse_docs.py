#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Jul 16 14:28:28 2020

@author: jmr
"""
import os
from glob import glob
from docx import Document
from docx2python import docx2python
from random import choice
import re
import subprocess
import json
import os
from os import listdir
from os.path import isfile, join
import shutil
import zipfile
import pandas as pd

## dirctories
os.chdir('/home/jmr/Dropbox/Current projects/thesis_papers/transparency, media, and compliance with HR Rulings/ecthr_media&compliance/data/case_docs_data')

### function for turning the ruling into text
files = glob("./rulings_dir/*")

## function for cleaning up the docx xml so that it can be read by python-docx
def update_docx(docname):
    """Update a docx such that it can be read by docx library.

        MSWord documents are a zip folder containing several XML files.
        As docx library cannot read 'smartTag', it is required to remove them.
        To do so, we open the zip, access the main XML file and manually sanitize it.

        :param docname: path to the document
        :type docname: str
        :return: path to the new document
        :rtype: str
    """
    # Remove temporary folder and files
    TMP = './tmp/echr_tmp_doc'
    try:
        shutil.rmtree(TMP)
    except:
        pass

    try:
        os.rm('./_proxy.docx')
    except:
        pass
    # Extract the document
    zip_ref = zipfile.ZipFile(docname, 'r')
    zip_ref.extractall(TMP)
    zip_ref.close()
    # Sanitize
    with open(os.path.join(TMP, 'word/document.xml'), 'r') as file:
        content = file.read()
        lines = content.split('>')
        remove_open = True
        for i, l in enumerate(lines):
            if '<w:smartTag ' in l and remove_open:
                del lines[i]
                remove_open = False
            if '</w:smartTag'==l and not remove_open:
                del lines[i]
                remove_open = True
        file.close()
    content = '>'.join(lines)
    # Recompress the archive
    with open(os.path.join(TMP, 'word/document.xml'), 'w') as file:
        file.write(content)
    shutil.make_archive('./proxy', 'zip', TMP)

    output_file = './_proxy.docx'
    os.rename('./proxy.zip', output_file)
    try:
        os.rm('./_proxy.docx')
    except:
        pass
    return output_file


### Generate the JSON file
# containers
dta_list = []
par_dict = {}
## start the loop
for current_file in files:
    ## parse the word document
    ## assign metadata
    par_dict['file'] = current_file.replace("./rulings_dir/", "")
    par_dict['doc_type'] = re.search("(?<=[A-Z]\_).+?(?=\_[0-9])", current_file).group(0)
    par_dict['doc_lang'] = re.search("(?<=rulings_dir/).+?(?=\_[A-Z])", current_file).group(0)
    par_dict['case_id'] = re.search("[0-9]+\_[0-9]+", current_file).group(0).replace('_', '/')    
    try:
        try:
            try:
                ## sanitize
                new_path = update_docx(current_file)
                ## first attempt
                doc = Document(new_path)
                par_dict['text_paragraphs'] = [cur_par.text for cur_par in doc.paragraphs]
            except:
                # failing, convert to docx using soffice
                subprocess.call(['soffice', '--headless', '--convert-to', 'docx', new_path])      
                doc = Document(new_path)
                par_dict['text_paragraphs'] = doc.body
        except:
            ## try with docx2python
            # parse and extract the paragraphs
            doc = docx2python(new_path)
            par_dict['text_paragraphs'] = doc.body
        ## export
        filename =  "./rulings_data/json/" + current_file.replace("./rulings_dir/", "").replace("docx", "json")
        with open(filename, 'w', encoding='utf8') as fout:
            json.dump(par_dict , fout, ensure_ascii=False)
        print(par_dict)
    except:
        print("couldn't parse %s"%current_file)

    
    